"""
Export Service — converts paper draft to Word (.docx) and PDF formats.
"""
import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table as RLTable, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.colors import HexColor

# Register Chinese font for PDF
try:
    pdfmetrics.registerFont(TTFont('STHeiti', '/System/Library/Fonts/STHeiti Medium.ttc', subfontIndex=0))
    CHINESE_FONT = 'STHeiti'
except Exception:
    CHINESE_FONT = 'Helvetica'


def _md_to_plain_text(md_text: str) -> str:
    """Strip markdown syntax for plain-text export (non-table content)."""
    text = md_text
    # Code blocks
    text = re.sub(r'```json\n(.*?)\n```', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'```\n?(.*?)\n?```', r'\1', text, flags=re.DOTALL)
    # Images
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    # Links
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # Bold/italic
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    # Headings
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    # Lists
    text = re.sub(r'^\s*[-*+]\s+', '  • ', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*\d+\.\s+', '  ', text, flags=re.MULTILINE)
    # Horizontal rules
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    # Extra newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _parse_md_tables(md_text: str) -> list[tuple[list[list[str]], int, int]]:
    """
    Extract markdown tables from text.
    Returns list of (rows, num_cols) tuples with their character positions.
    """
    tables = []
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        if '|' in lines[i] and re.match(r'\s*\|.*\|\s*', lines[i]):
            # Found table start
            rows = []
            # Collect header row
            header = [c.strip() for c in lines[i].strip('|').split('|')]
            rows.append(header)
            i += 1
            # Skip separator row (e.g., |---|---|)
            if i < len(lines) and re.match(r'\s*\|[\s\-:]+\|\s*', lines[i]):
                i += 1
            # Collect data rows
            while i < len(lines) and '|' in lines[i] and re.match(r'\s*\|.*\|\s*', lines[i]):
                row = [c.strip() for c in lines[i].strip('|').split('|')]
                rows.append(row)
                i += 1
            if len(rows) >= 2:
                num_cols = len(rows[0])
                tables.append((rows, num_cols))
            continue
        i += 1
    return tables


def _split_sections(md_text: str) -> list[tuple[str, str]]:
    """Split markdown text into (heading, body) sections."""
    sections = []
    current_heading = ""
    current_body = ""
    heading_pattern = re.compile(r'^#{1,6}\s+(.+)$', re.MULTILINE)

    for line in md_text.split('\n'):
        m = heading_pattern.match(line)
        if m:
            if current_heading or current_body.strip():
                sections.append((current_heading, current_body.strip()))
            current_heading = m.group(1).strip()
            current_body = ""
        else:
            current_body += line + "\n"

    if current_heading or current_body.strip():
        sections.append((current_heading, current_body.strip()))

    return sections


def _render_content(doc: Document, body: str):
    """
    Render markdown body content to docx, handling tables properly.
    Tables are extracted and rendered as docx Table objects.
    Remaining text is processed with _md_to_plain_text.
    """
    lines = body.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        # Check if this is a table row
        if '|' in line and re.match(r'\s*\|.*\|\s*', line):
            # Collect all consecutive table rows
            table_lines = []
            while i < len(lines) and '|' in lines[i] and re.match(r'\s*\|.*\|\s*', lines[i]):
                table_lines.append(lines[i])
                i += 1
            # Skip separator row
            if i < len(lines) and re.match(r'\s*\|[\s\-:]+\|\s*', lines[i]):
                i += 1

            # Parse into rows
            rows = []
            for tl in table_lines:
                if not re.match(r'\s*\|[\s\-:]+\|\s*', tl):
                    row = [c.strip() for c in tl.strip('|').split('|')]
                    rows.append(row)

            if len(rows) >= 2 and len(rows[0]) >= 1:
                num_cols = len(rows[0])
                tbl = doc.add_table(rows=len(rows), cols=num_cols)
                tbl.style = 'Table Grid'
                for r_idx, row_data in enumerate(rows):
                    for c_idx, cell_text in enumerate(row_data):
                        cell = tbl.rows[r_idx].cells[c_idx]
                        # Strip markdown from cell text
                        clean = _md_to_plain_text(cell_text)
                        cell.text = clean
                continue
        else:
            # Regular paragraph
            stripped = _md_to_plain_text(line)
            if stripped:
                doc.add_paragraph(stripped, style='Normal')
        i += 1


def export_word(md_content: str, title: str = "AI Paper") -> bytes:
    """Convert markdown paper to .docx bytes with proper table support."""
    doc = Document()
    doc.core_properties.title = title

    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    sections = _split_sections(md_content)

    for section_title, section_body in sections:
        if section_title:
            doc.add_heading(section_title, level=1)
        if section_body:
            _render_content(doc, section_body)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(md_content: str, title: str = "AI Paper") -> bytes:
    """Convert markdown paper to PDF bytes using reportlab with Chinese font support."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5*cm,
        rightMargin=2.5*cm,
        topMargin=2.5*cm,
        bottomMargin=2.5*cm,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'TitleStyle',
        fontName=CHINESE_FONT,
        fontSize=18,
        leading=22,
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    heading_style = ParagraphStyle(
        'HeadingStyle',
        fontName=CHINESE_FONT,
        fontSize=14,
        leading=18,
        spaceBefore=16,
        spaceAfter=6,
        textColor=HexColor('#1a1a1a'),
    )
    body_style = ParagraphStyle(
        'BodyStyle',
        fontName=CHINESE_FONT,
        fontSize=11,
        leading=16,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
    )
    table_header_style = ParagraphStyle(
        'TableHeader',
        fontName=CHINESE_FONT,
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
    )
    table_cell_style = ParagraphStyle(
        'TableCell',
        fontName=CHINESE_FONT,
        fontSize=10,
        leading=14,
        alignment=TA_LEFT,
    )

    story = []

    story.append(Paragraph(title, title_style))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body_style))
    story.append(Spacer(1, 0.5*cm))

    sections = _split_sections(md_content)

    for section_title, section_body in sections:
        if section_title:
            story.append(Paragraph(section_title, heading_style))
        if section_body:
            _render_pdf_content(story, section_body, body_style, table_header_style, table_cell_style)
        story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    return buf.getvalue()


def _render_pdf_content(story: list, body: str, body_style: ParagraphStyle,
                       table_header_style: ParagraphStyle, table_cell_style: ParagraphStyle):
    """Render markdown body to PDF story, handling tables."""
    lines = body.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        # Check if this is a table row
        if '|' in line and re.match(r'\s*\|.*\|\s*', line):
            # Collect all consecutive table rows
            table_lines = []
            while i < len(lines) and '|' in lines[i] and re.match(r'\s*\|.*\|\s*', lines[i]):
                table_lines.append(lines[i])
                i += 1
            # Skip separator row
            if i < len(lines) and re.match(r'\s*\|[\s\-:]+\|\s*', lines[i]):
                i += 1

            rows = []
            for tl in table_lines:
                if not re.match(r'\s*\|[\s\-:]+\|\s*', tl):
                    row = [c.strip() for c in tl.strip('|').split('|')]
                    rows.append(row)

            if len(rows) >= 2 and len(rows[0]) >= 1:
                num_cols = len(rows[0])
                # Build table data for reportlab
                rl_rows = []
                for r_idx, row_data in enumerate(rows):
                    rl_row = []
                    for c_idx, cell_text in enumerate(row_data):
                        clean = _md_to_plain_text(cell_text)
                        style = table_header_style if r_idx == 0 else table_cell_style
                        rl_row.append(Paragraph(clean, style))
                    rl_rows.append(rl_row)

                col_width = (A4[0] - 5*cm) / num_cols
                rl_table = RLTable(rl_rows, colWidths=[col_width] * num_cols)
                rl_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), HexColor('#e8e8e8')),
                    ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#cccccc')),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ]))
                story.append(rl_table)
                story.append(Spacer(1, 0.3*cm))
            continue
        else:
            stripped = _md_to_plain_text(line)
            if stripped:
                story.append(Paragraph(stripped, body_style))
        i += 1
